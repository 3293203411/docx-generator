from fastapi import FastAPI, HTTPException
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from docx import Document
import requests
import io
import uuid
import tempfile
import os

app = FastAPI()

# 配置CORS，允许跨域请求
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# 全局内存存储（Render免费版重启会清空）
file_store = {}

# 健康检查接口
@app.get("/health")
async def health_check():
    return {"status": "healthy"}

# 文档生成接口
@app.post("/generate")
async def generate_document(data: dict):
    try:
        # 接收参数
        text_keys = data.get("text_keys", [])
        text_values = data.get("text_values", [])
        template_file_url = data.get("template_file_url", "")
        
        if not template_file_url:
            raise HTTPException(status_code=400, detail="template_file_url不能为空")
        
        if len(text_keys) != len(text_values):
            raise HTTPException(status_code=400, detail=f"text_keys数量({len(text_keys)})与text_values数量({len(text_values)})不匹配")
        
        # 1. 下载模板文件
        try:
            template_response = requests.get(template_file_url, timeout=30)
            template_response.raise_for_status()
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"模板文件下载失败: {str(e)}")
        
        # 2. 打开Word模板
        try:
            template_stream = io.BytesIO(template_response.content)
            doc = Document(template_stream)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"模板文件打开失败: {str(e)}")
        
        # 3. 替换普通段落中的占位符（保留原有格式）
        for paragraph in doc.paragraphs:
            for key, value in zip(text_keys, text_values):
                placeholder = f"{{{{{key}}}}}"
                if placeholder in paragraph.text:
                    # 遍历每个run，只替换占位符内容，保留原有格式（字体、下划线、颜色等）
                    for run in paragraph.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, str(value))
        
        # 4. 替换表格中的占位符（保留表格内的格式，很多用户模板里有表格）
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in zip(text_keys, text_values):
                            placeholder = f"{{{{{key}}}}}"
                            if placeholder in paragraph.text:
                                # 同样保留表格内的原有格式
                                for run in paragraph.runs:
                                    if placeholder in run.text:
                                        run.text = run.text.replace(placeholder, str(value))
        
        # 5. 保存生成的文档到内存
        output_stream = io.BytesIO()
        doc.save(output_stream)
        output_stream.seek(0)
        
        # 6. 生成唯一文件名（避免中文编码问题）
        file_id = str(uuid.uuid4())[:8]
        filename = f"立项文档_{file_id}.docx"
        
        # 7. 存储到内存
        file_store[file_id] = {
            "content": output_stream.getvalue(),
            "filename": filename
        }
        
        # 8. 返回下载链接
        base_url = str(requests.get('https://api.ipify.org').text)  # 自动获取当前服务域名
        download_url = f"https://docx-generator-2eox.onrender.com/download/{file_id}"
        
        return {
            "success": True,
            "full_download_url": download_url,
            "filename": filename,
            "replaced_count": len(text_keys)
        }
    
    except HTTPException as e:
        raise e
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"内部错误: {str(e)}")

# 文件下载接口
@app.get("/download/{file_id}")
async def download_file(file_id: str):
    if file_id not in file_store:
        raise HTTPException(status_code=404, detail="文件不存在或已过期")
    
    file_data = file_store[file_id]
    
    # 生成临时文件返回
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(file_data["content"])
        tmp_path = tmp.name
    
    return FileResponse(
        tmp_path,
        filename=file_data["filename"],
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=int(os.environ.get("PORT", 8000)))
